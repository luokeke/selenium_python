#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time : 2020/1/8 10:41
# @Author : liuhuiling

from docx import Document
document  = Document()
document.add_heading("Document Title",0)
p = document.add_paragraph('A plain paragarph having some')
document.add_page_break()
document.save('demo.docx')

#http://python-docx.readthedocs.io/en/latest/index.html